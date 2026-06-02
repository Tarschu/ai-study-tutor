#!/usr/bin/env python3
"""Import course materials into generated reference notes."""

from __future__ import annotations

import argparse
import re
from datetime import datetime
from pathlib import Path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Extract text from course materials and create generated reference notes."
    )
    parser.add_argument("files", nargs="+", help="PDF, DOCX, TXT, or Markdown files")
    parser.add_argument("--course", required=True, help="Course name")
    parser.add_argument("--topic", default="", help="Chapter or knowledge point")
    parser.add_argument(
        "--output-dir",
        default=None,
        help="Output directory. Defaults to the plugin generated references folder.",
    )
    parser.add_argument(
        "--max-chars",
        type=int,
        default=60000,
        help="Maximum extracted characters to keep per imported note.",
    )
    return parser.parse_args()


def plugin_generated_dir() -> Path:
    return (
        Path(__file__).resolve().parents[1]
        / "skills"
        / "ai-study-tutor"
        / "references"
        / "generated"
    )


def slugify(value: str) -> str:
    value = value.strip().lower()
    value = re.sub(r"[^\w\u4e00-\u9fff]+", "-", value)
    value = re.sub(r"-+", "-", value).strip("-")
    return value or "course-material"


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".markdown"}:
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".docx":
        return extract_docx(path)
    raise SystemExit(f"Unsupported file type: {path}")


def extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ModuleNotFoundError as exc:
        raise SystemExit("PDF import requires pypdf in the active Python environment.") from exc

    reader = PdfReader(str(path))
    chunks: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        chunks.append(f"\n\n--- Page {index} ---\n{text.strip()}")
    return "\n".join(chunks).strip()


def extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ModuleNotFoundError as exc:
        raise SystemExit("DOCX import requires python-docx in the active Python environment.") from exc

    doc = Document(str(path))
    paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    table_rows: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            table_rows.append(" | ".join(cells))
    return "\n".join(paragraphs + table_rows).strip()


def first_nonempty_lines(text: str, limit: int = 20) -> list[str]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return lines[:limit]


def write_import_note(
    source: Path,
    text: str,
    *,
    course: str,
    topic: str,
    output_dir: Path,
    max_chars: int,
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    today = datetime.now().strftime("%Y-%m-%d")
    topic_part = slugify(topic) if topic else "general"
    output = output_dir / f"{slugify(course)}-{topic_part}-{slugify(source.stem)}.md"
    excerpt = text[:max_chars].strip()
    if len(text) > max_chars:
        excerpt += "\n\n[Truncated by import_course_material.py]"
    preview = "\n".join(f"- {line}" for line in first_nonempty_lines(text, 12))
    content = f"""# 课程资料导入：{source.name}

## Metadata

- Course: {course}
- Topic: {topic or "未指定"}
- Source: {source}
- Imported: {today}
- Characters: {len(text)}

## Quick Preview

{preview or "- 未提取到可读文本"}

## How To Use

- Use this generated note only as course context.
- When solving a problem, still verify the actual user-provided question first.
- Search this file for formulas, definitions, theorem names, examples, or chapter terms.

## Extracted Text

{excerpt or "未提取到可读文本。若这是扫描版 PDF 或图片，请先使用图片预处理/OCR 工具。"}
"""
    output.write_text(content, encoding="utf-8")
    return output


def main() -> None:
    args = parse_args()
    output_dir = Path(args.output_dir).expanduser().resolve() if args.output_dir else plugin_generated_dir()
    outputs: list[Path] = []
    for raw in args.files:
        source = Path(raw).expanduser().resolve()
        if not source.is_file():
            raise SystemExit(f"Input file does not exist: {source}")
        text = extract_text(source)
        outputs.append(
            write_import_note(
                source,
                text,
                course=args.course,
                topic=args.topic,
                output_dir=output_dir,
                max_chars=args.max_chars,
            )
        )
    print("Generated:")
    for output in outputs:
        print(output)


if __name__ == "__main__":
    main()
